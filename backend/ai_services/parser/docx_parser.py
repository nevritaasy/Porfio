from __future__ import annotations

import sys
from pathlib import Path

try:
    from docx import Document
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False


def extract_text_from_docx(docx_path: str | Path) -> dict:
    docx_path = Path(docx_path)
    if not docx_path.is_file():
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")

    if not _DOCX_AVAILABLE:
        raise ImportError("python-docx is not installed. Run: pip install python-docx")

    paragraphs: list[str] = []

    try:
        doc = Document(str(docx_path))
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_texts:
                    paragraphs.append("  |  ".join(row_texts))

    except Exception as exc:
        print(f"[docx_parser] Failed to parse DOCX: {exc}", file=sys.stderr)
        return {
            "text": "",
            "metadata": {
                "total_pages": 1,
                "quality": "poor",
                "extraction_method": "docx",
            },
        }

    full_text = "\n".join(paragraphs)
    quality = "good" if len(full_text) > 200 else ("partial" if len(full_text) > 50 else "poor")

    return {
        "text": full_text,
        "metadata": {
            "total_pages": 1,
            "quality": quality,
            "extraction_method": "docx",
        },
    }
